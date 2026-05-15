from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from docx import Document


class Word2MdProjectLayoutTests(unittest.TestCase):
    def test_converter_imports_from_project_package_and_writes_markdown(self) -> None:
        from word2md.converter import convert

        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            docx_path = root / "sample.docx"
            output_dir = root / "out"

            document = Document()
            document.add_heading("Sample Title", level=1)
            document.add_paragraph("Hello from Word2MD")
            document.save(docx_path)

            result = convert(docx_path, output_dir=output_dir)

            self.assertEqual(output_dir / "sample.md", result.output_path)
            self.assertTrue(result.output_path.exists())
            self.assertIn("# Sample Title", result.markdown)
            self.assertIn("Hello from Word2MD", result.markdown)

    def test_gui_entrypoint_imports_from_project_package(self) -> None:
        from word2md.gui import main

        self.assertTrue(callable(main))


if __name__ == "__main__":
    unittest.main()
